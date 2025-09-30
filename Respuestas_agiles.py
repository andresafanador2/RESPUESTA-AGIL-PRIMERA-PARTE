import tkinter as tk
from tkinter import ttk
from docx import Document
import os

# Diccionario de grupos y subgrupos con el formato que deseas (código - nombre)
grupos = {
    "01": {
        "0111 - Cultivo de cereales (excepto arroz), legumbres y semillas oleaginosas",
        "0112 - Cultivo de arroz",
        "0113 - Cultivo de hortalizas, raíces y tubérculos",
        "0114 - Cultivo de tabaco",
        "0115 - Cultivo de plantas textiles",
        "0119 - Otros cultivos transitorios n.c.p.",
        "0121 - Cultivo de frutas tropicales y subtropicales",
        "0122 - Cultivo de plátano y banano",
        "0123 - Cultivo de café",
        "0124 - Cultivo de caña de azúcar",
        "0125 - Cultivo de flor de corte",
        "0126 - Cultivo de palma para aceite (palma africana) y otros frutos oleaginosos",
        "0127 - Cultivo de plantas con las que se preparan bebidas",
        "0128 - Cultivo de especias y de plantas aromáticas y medicinales",
        "0129 - Otros cultivos permanentes n.c.p.",
        "0130 - Propagación de plantas (actividades de los viveros, excepto viveros forestales)",
        "0141 - Cría de ganado bovino y bufalino",
        "0142 - Cría de caballos y otros equinos",
        "0143 - Cría de ovejas y cabras",
        "0144 - Cría de ganado porcino",
        "0145 - Cría de aves de corral",
        "0149 - Cría de otros animales n.c.p.",
        "0150 - Explotación mixta (agrícola y pecuaria)",
        "0161 - Actividades de apoyo a la agricultura",
        "0162 - Actividades de apoyo a la ganadería",
        "0163 - Actividades posteriores a la cosecha",
        "0164 - Tratamiento de semillas para propagación",
        "0170 - Caza ordinaria y mediante trampas y actividades de servicios conexas"
    },
      "02": {
    "0210 - Silvicultura y otras actividades forestales",
    "0220 - Extracción de madera",
    "0230 - Recolección de productos forestales diferentes a la madera",
    "0240 - Servicios de apoyo a la silvicultura"
  },
  "03": {
    "0311 - Pesca marítima",
    "0312 - Pesca de agua dulce",
    "0321 - Acuicultura marítima",
    "0322 - Acuicultura de agua dulce"
  },
  "05": {
    "0510 - Extracción de hulla (carbón de piedra)",
    "0520 - Extracción de carbón lignito"
  },
  "06": {
    "0610 - Extracción de petróleo crudo",
    "0620 - Extracción de gas natural"
  },
  "07": {
    "0710 - Extracción de minerales de hierro",
    "0721 - Extracción de minerales de uranio y de torio",
    "0722 - Extracción de oro y otros metales preciosos",
    "0723 - Extracción de minerales de níquel",
    "0729 - Extracción de otros minerales metalíferos no ferrosos n.c.p."
  },
  "08": {
    "0811 - Extracción de piedra, arena, arcillas comunes, yeso y anhidrita",
    "0812 - Extracción de arcillas de uso industrial, caliza, caolín y bentonitas",
    "0820 - Extracción de esmeraldas, piedras preciosas y semipreciosas",
    "0891 - Extracción de minerales para la fabricación de abonos y productos químicos",
    "0892 - Extracción de halita (sal)",
    "0899 - Extracción de otros minerales no metálicos n.c.p."
  },
  "09": {
    "0910 - Actividades de apoyo para la extracción de petróleo y de gas natural",
    "0990 - Actividades de apoyo para otras actividades de explotación de minas y canteras"
  },
  "10": {
    "1011 - Procesamiento y conservación de carne y productos cárnicos",
    "1012 - Procesamiento y conservación de pescados, crustáceos y moluscos",
    "1020 - Procesamiento y conservación de frutas, legumbres, hortalizas y tubérculos",
    "1031 - Extracción de aceites de origen vegetal crudos",
    "1032 - Elaboración de aceites y grasas de origen vegetal refinados",
    "1033 - Elaboración de aceites y grasas de origen animal",
    "1040 - Elaboración de productos lácteos",
    "1051 - Elaboración de productos de molinería",
    "1052 - Elaboración de almidones y productos derivados del almidón",
    "1061 - Trilla de café",
    "1062 - Descafeinado, tostión y molienda del café",
    "1063 - Otros derivados del café",
    "1071 - Elaboración y refinación de azúcar",
    "1072 - Elaboración de panela",
    "1081 - Elaboración de productos de panadería",
    "1082 - Elaboración de cacao, chocolate y productos de confitería",
    "1083 - Elaboración de macarrones, fideos, alcuzcuz y productos farináceos similares",
    "1084 - Elaboración de comidas y platos preparados",
    "1089 - Elaboración de otros productos alimenticios n.c.p.",
    "1090 - Elaboración de alimentos preparados para animales"
  },
  "11": {
    "1101 - Destilación, rectificación y mezcla de bebidas alcohólicas",
    "1102 - Elaboración de bebidas fermentadas no destiladas",
    "1103 - Producción de malta, elaboración de cervezas y otras bebidas malteadas",
    "1104 - Elaboración de bebidas no alcohólicas, producción de aguas minerales y otras aguas embotelladas"
  },
  "12": {
    "1200 - Elaboración de productos de tabaco"
  },
  "13": {
    "1311 - Preparación e hilatura de fibras textiles",
    "1312 - Tejeduría de productos textiles",
    "1313 - Acabado de productos textiles",
    "1391 - Fabricación de tejidos de punto y ganchillo",
    "1392 - Confección de artículos con materiales textiles, excepto prendas de vestir",
    "1393 - Fabricación de tapetes y alfombras para pisos",
    "1394 - Fabricación de cuerdas, cordeles, cables, bramantes y redes",
    "1399 - Fabricación de otros artículos textiles n.c.p."
  },
  "14": {
    "1410 - Confección de prendas de vestir, excepto prendas de piel",
    "1420 - Fabricación de artículos de piel",
    "1430 - Fabricación de artículos de punto y ganchillo"
  },
  "15": {
    "1511 - Curtido y recurtido de cueros; recurtido y teñido de pieles",
    "1512 - Fabricación de artículos de viaje, bolsos de mano y artículos similares elaborados en cuero, y fabricación de artículos de talabartería y guarnicionería",
    "1513 - Fabricación de artículos de viaje, bolsos de mano y artículos similares; artículos de talabartería y guarnicionería elaborados en otros materiales",
    "1521 - Fabricación de calzado de cuero y piel, con cualquier tipo de suela",
    "1522 - Fabricación de otros tipos de calzado, excepto calzado de cuero y piel",
    "1523 - Fabricación de partes del calzado"
  },
  "16": {
    "1610 - Aserrado, acepillado e impregnación de la madera",
    "1620 - Fabricación de hojas de madera para enchapado; fabricación de tableros contrachapados, tableros laminados, tableros de partículas y otros tableros y paneles",
    "1630 - Fabricación de partes y piezas de madera, de carpintería y ebanistería para la construcción",
    "1640 - Fabricación de recipientes de madera",
    "1690 - Fabricación de otros productos de madera; fabricación de artículos de corcho, cestería y espartería"
  },
  "17": {
    "1701 - Fabricación de pulpas (pastas) celulósicas; papel y cartón",
    "1702 - Fabricación de papel y cartón ondulado (corrugado); fabricación de envases, empaques y de embalajes de papel y cartón",
    "1709 - Fabricación de otros artículos de papel y cartón"
  },
  "18": {
    "1811 - Actividades de impresión",
    "1812 - Actividades de servicios relacionados con la impresión",
    "1820 - Producción de copias a partir de grabaciones originales"
  },
  "19": {
    "1910 - Fabricación de productos de hornos de coque",
    "1921 - Fabricación de productos de la refinación del petróleo",
    "1922 - Actividad de mezcla de combustibles"
  },
  "20": {
    "2011 - Fabricación de sustancias y productos químicos básicos",
    "2012 - Fabricación de abonos y compuestos inorgánicos nitrogenados",
    "2013 - Fabricación de plásticos en formas primarias",
    "2014 - Fabricación de caucho sintético en formas primarias",
    "2021 - Fabricación de plaguicidas y otros productos químicos de uso agropecuario",
    "2022 - Fabricación de pinturas, barnices y revestimientos similares, tintas para impresión y masillas",
    "2023 - Fabricación de jabones y detergentes, preparados para limpiar y pulir; perfumes y preparados de tocador",
    "2029 - Fabricación de otros productos químicos n.c.p.",
    "2030 - Fabricación de fibras sintéticas y artificiales"
  },
  "21": {
    "2100 - Fabricación de productos farmacéuticos, sustancias químicas medicinales y productos botánicos de uso farmacéutico"
  },
  "22": {
    "2211 - Fabricación de llantas y neumáticos de caucho",
    "2212 - Reencauche de llantas usadas",
    "2219 - Fabricación de formas básicas de caucho y otros productos de caucho n.c.p.",
    "2221 - Fabricación de formas básicas de plástico",
    "2229 - Fabricación de artículos de plástico n.c.p."
  },
  "23": {
    "2310 - Fabricación de vidrio y productos de vidrio",
    "2391 - Fabricación de productos refractarios",
    "2392 - Fabricación de materiales de arcilla para la construcción",
    "2393 - Fabricación de otros productos de cerámica y porcelana",
    "2394 - Fabricación de cemento, cal y yeso",
    "2395 - Fabricación de artículos de hormigón, cemento y yeso",
    "2396 - Corte, tallado y acabado de la piedra",
    "2399 - Fabricación de otros productos minerales no metálicos n.c.p."
  },
  "24": {
    "2410 - Industrias básicas de hierro y de acero",
    "2421 - Industrias básicas de metales preciosos",
    "2429 - Industrias básicas de otros metales no ferrosos",
    "2431 - Fundición de hierro y de acero",
    "2432 - Fundición de metales no ferrosos"
  },
  "25": {
    "2511 - Fabricación de productos metálicos para uso estructural",
    "2512 - Fabricación de tanques, depósitos y recipientes de metal, excepto los utilizados para el envase o transporte de mercancías",
    "2513 - Fabricación de generadores de vapor, excepto calderas de agua caliente para calefacción central",
    "2520 - Fabricación de armas y municiones",
    "2591 - Forja, prensado, estampado y laminado de metal; pulvimetalurgia",
    "2592 - Tratamiento y revestimiento de metales; mecanizado",
    "2593 - Fabricación de artículos de cuchillería, herramientas de mano y artículos de ferretería",
    "2599 - Fabricación de otros productos elaborados de metal n.c.p."
  },
  "26": {
    "2610 - Fabricación de componentes y tableros electrónicos",
    "2620 - Fabricación de computadoras y de equipo periférico",
    "2630 - Fabricación de equipos de comunicación",
    "2640 - Fabricación de aparatos electrónicos de consumo",
    "2651 - Fabricación de equipo de medición, prueba, navegación y control",
    "2652 - Fabricación de relojes",
    "2660 - Fabricación de equipo de irradiación y equipo electrónico de uso médico y terapéutico",
    "2670 - Fabricación de instrumentos ópticos y equipo fotográfico",
    "2680 - Fabricación de medios magnéticos y ópticos para almacenamiento de datos"
  },
  "27": {
    "2711 - Fabricación de motores, generadores y transformadores eléctricos",
    "2712 - Fabricación de aparatos de distribución y control de la energía eléctrica",
    "2720 - Fabricación de pilas, baterías y acumuladores eléctricos",
    "2731 - Fabricación de hilos y cables eléctricos y de fibra óptica",
    "2732 - Fabricación de dispositivos de cableado",
    "2740 - Fabricación de equipos eléctricos de iluminación",
    "2750 - Fabricación de aparatos de uso doméstico",
    "2790 - Fabricación de otros tipos de equipo eléctrico n.c.p."
  },
  "28": {
    "2811 - Fabricación de motores, turbinas, y partes para motores de combustión interna",
    "2812 - Fabricación de equipos de potencia hidráulica y neumática",
    "2813 - Fabricación de otras bombas, compresores, grifos y válvulas",
    "2814 - Fabricación de cojinetes, engranajes, trenes de engranajes y piezas de transmisión",
    "2815 - Fabricación de hornos, hogares y quemadores industriales",
    "2816 - Fabricación de equipo de elevación y manipulación",
    "2817 - Fabricación de maquinaria y equipo de oficina (excepto computadoras y equipo periférico)",
    "2818 - Fabricación de herramientas manuales con motor",
    "2819 - Fabricación de otros tipos de maquinaria y equipo de uso general n.c.p.",
    "2821 - Fabricación de maquinaria agropecuaria y forestal",
    "2822 - Fabricación de máquinas formadoras de metal y de máquinas herramienta",
    "2823 - Fabricación de maquinaria para la metalurgia",
    "2824 - Fabricación de maquinaria para explotación de minas y canteras y para obras de construcción",
    "2825 - Fabricación de maquinaria para la elaboración de alimentos, bebidas y tabaco",
    "2826 - Fabricación de maquinaria para la elaboración de productos textiles, prendas de vestir y cueros",
    "2829 - Fabricación de otros tipos de maquinaria y equipo de uso especial n.c.p."
  },
  "29": {
    "2910 - Fabricación de vehículos automotores y sus motores",
    "2920 - Fabricación de carrocerías para vehículos automotores; fabricación de remolques y semirremolques",
    "2930 - Fabricación de partes, piezas (autopartes) y accesorios (lujos) para vehículos automotores"
  },
  "30": {
    "3011 - Construcción de barcos y de estructuras flotantes",
    "3012 - Construcción de embarcaciones de recreo y deporte",
    "3020 - Fabricación de locomotoras y de material rodante para ferrocarriles",
    "3030 - Fabricación de aeronaves, naves espaciales y de maquinaria conexa",
    "3040 - Fabricación de vehículos militares de combate",
    "3091 - Fabricación de motocicletas",
    "3092 - Fabricación de bicicletas y de sillas de ruedas para personas con discapacidad",
    "3099 - Fabricación de otros tipos de equipo de transporte n.c.p."
  },
  "31": {
    "3110 - Fabricación de muebles",
    "3120 - Fabricación de colchones y somieres"
  },
  "32": {
    "3211 - Fabricación de joyas y artículos conexos",
    "3212 - Fabricación de bisutería y artículos conexos",
    "3220 - Fabricación de instrumentos musicales",
    "3230 - Fabricación de artículos y equipo para la práctica del deporte",
    "3240 - Fabricación de juegos, juguetes y rompecabezas",
    "3250 - Fabricación de instrumentos, aparatos y materiales médicos y odontológicos (incluido mobiliario)",
    "3290 - Otras industrias manufactureras n.c.p."
  },
  "33": {
    "3311 - Mantenimiento y reparación especializado de productos elaborados en metal",
    "3312 - Mantenimiento y reparación especializado de maquinaria y equipo",
    "3313 - Mantenimiento y reparación especializado de equipo electrónico y óptico",
    "3314 - Mantenimiento y reparación especializado de equipo eléctrico",
    "3315 - Mantenimiento y reparación especializado de equipo de transporte, excepto los vehículos automotores, motocicletas y bicicletas",
    "3319 - Mantenimiento y reparación de otros tipos de equipos y sus componentes n.c.p.",
    "3320 - Instalación especializada de maquinaria y equipo industrial"
  },
  "35": {
    "3511 - Generación de energía eléctrica",
    "3512 - Transmisión de energía eléctrica",
    "3513 - Distribución de energía eléctrica",
    "3514 - Comercialización de energía eléctrica",
    "3520 - Producción de gas; distribución de combustibles gaseosos por tuberías",
    "3530 - Suministro de vapor y aire acondicionado"
  },
  "36": {
    "3600 - Captación, tratamiento y distribución de agua"
  },
  "41": {
    "4111 - Construcción de edificios residenciales",
    "4112 - Construcción de edificios no residenciales"
  },
  "42": {
    "4210 - Construcción de carreteras y vías de ferrocarril",
    "4220 - Construcción de proyectos de servicio público",
    "4290 - Construcción de otras obras de ingeniería civil"
  },
  "43": {
    "4311 - Demolición",
    "4312 - Preparación del terreno",
    "4321 - Instalaciones eléctricas",
    "4322 - Instalaciones de fontanería, calefacción y aire acondicionado",
    "4329 - Otras instalaciones especializadas",
    "4330 - Terminación y acabado de edificios y obras de ingeniería civil"
  },
}

# Lista de municipios de Norte de Santander
municipios_norte_santander = [
    "Arboledas", "Cucutilla", "Gramalote", "Lourdes", "Salazar de Las Palmas",
    "Santiago", "Villa Caro", "Cúcuta", "El Zulia", "Los Patios",
    "Puerto Santander", "San Cayetano", "Villa del Rosario", "Bucarasica",
    "El Tarra", "Sardinata", "Tibú", "Ábrego", "Cáchira", "Convención",
    "El Carmen", "Hacarí", "La Esperanza", "La Playa de Belén", "Ocaña",
    "San Calixto", "Teorama", "Cácota", "Chitagá", "Mutiscua",
    "Pamplona", "Pamplonita", "Santo Domingo de Silos", "Bochalema",
    "Chinácota", "Durania", "Herrán", "Labateca", "Ragonvalia", "Toledo"
]

# Variables globales para los campos del formulario
nombre_usuario = ""
codigo_cliente = ""
direccion = ""
correo = ""
telefono = ""
municipio = ""
radicado = ""
fecha_radicado = ""
expediente = ""
proceso = ""
grupo_seleccionado = ""
subgrupo_seleccionado = ""
subgrupo_codigo = ""  # Variable para almacenar el código del subgrupo
subgrupo_nombre = ""  # Variable para almacenar el nombre del subgrupo

# Crear la ventana principal
ventana = tk.Tk()
ventana.title("Formulario de Selección y Datos")
ventana.geometry("600x800")

# Frame para los formularios
frame_formulario = tk.Frame(ventana)
frame_formulario.pack(pady=20)

# Frame para los desplegables (para Grupo y Subgrupo)
frame_desplegables = tk.Frame(ventana)
frame_desplegables.pack(pady=20)

# Función para mostrar formulario
def mostrar_formulario():
    global grupo_combo, subgrupo_combo 
    global nombre_usuario, codigo_cliente, direccion, correo, telefono, municipio, radicado, fecha_radicado, expediente, proceso, grupo_seleccionado, subgrupo_seleccionado
    
    seleccion = combo_box.get()

    # Limpiar el formulario previo antes de mostrar el nuevo
    for widget in frame_formulario.winfo_children():
        widget.destroy()

    # Limpiar los desplegables de Grupo y Subgrupo
    for widget in frame_desplegables.winfo_children():
        widget.destroy()

    # Si la opción es "Aprobación exención contribución" o "Negación exención contribución"
    if seleccion in ["Aprobación exención contribución", "Negación exención contribución"]:
        # Crear los widgets de grupo y subgrupo nuevamente
        tk.Label(frame_desplegables, text="Grupo").pack(pady=5)
        grupo_combo = ttk.Combobox(frame_desplegables, values=list(grupos.keys()), state="normal")
        grupo_combo.set("01")
        grupo_combo.bind("<<ComboboxSelected>>", actualizar_subgrupos)
        grupo_combo.pack(pady=5)

        tk.Label(frame_desplegables, text="Subgrupo (Número y Nombre)").pack(pady=5)
        subgrupo_combo = ttk.Combobox(frame_desplegables, state="normal", width=40)  # Aumento el ancho del ComboBox
        subgrupo_combo.pack(pady=5)
        
        # Llenar el formulario con los campos necesarios
        crear_campos_formulario()

    else:
        # Deshabilitar los desplegables y ponerlos en gris
        grupo_combo = ttk.Combobox(frame_desplegables, values=list(grupos.keys()), state="disabled")
        grupo_combo.set("Seleccione un grupo")
        grupo_combo.config(background="lightgray")
        grupo_combo.pack(pady=5)

        subgrupo_combo = ttk.Combobox(frame_desplegables, state="disabled", width=40)  # Aumento el ancho del ComboBox
        subgrupo_combo.config(background="lightgray")
        subgrupo_combo.pack(pady=5)

        # Solo mostrar el formulario sin los desplegables de grupo y subgrupo
        crear_campos_formulario()

def crear_campos_formulario():
    global nombre_usuario, codigo_cliente, direccion, correo, telefono, municipio, radicado, fecha_radicado, expediente, proceso
    
    # Llenar el formulario con los campos necesarios
    tk.Label(frame_formulario, text="Nombre de usuario").grid(row=0, column=0)
    nombre_usuario = tk.Entry(frame_formulario)
    nombre_usuario.grid(row=0, column=1)

    tk.Label(frame_formulario, text="Código de cliente").grid(row=1, column=0)
    codigo_cliente = tk.Entry(frame_formulario)
    codigo_cliente.grid(row=1, column=1)

    tk.Label(frame_formulario, text="Dirección").grid(row=2, column=0)
    direccion = tk.Entry(frame_formulario)
    direccion.grid(row=2, column=1)

    tk.Label(frame_formulario, text="Correo").grid(row=3, column=0)
    correo = tk.Entry(frame_formulario)
    correo.grid(row=3, column=1)

    tk.Label(frame_formulario, text="Teléfono").grid(row=4, column=0)
    telefono = tk.Entry(frame_formulario)
    telefono.grid(row=4, column=1)

    tk.Label(frame_formulario, text="Municipio").grid(row=5, column=0)
    municipio = tk.Entry(frame_formulario)
    municipio.grid(row=5, column=1)

    tk.Label(frame_formulario, text="Radicado").grid(row=6, column=0)
    radicado = tk.Entry(frame_formulario)
    radicado.grid(row=6, column=1)

    tk.Label(frame_formulario, text="Fecha de radicado").grid(row=7, column=0)
    fecha_radicado = tk.Entry(frame_formulario)
    fecha_radicado.grid(row=7, column=1)

    tk.Label(frame_formulario, text="Expediente").grid(row=8, column=0)
    expediente = tk.Entry(frame_formulario)
    expediente.grid(row=8, column=1)

    tk.Label(frame_formulario, text="Proceso").grid(row=9, column=0)
    proceso = tk.Entry(frame_formulario)
    proceso.grid(row=9, column=1)

# Función para actualizar subgrupos según el grupo seleccionado
def actualizar_subgrupos(event=None):
    global subgrupo_seleccionado, subgrupo_codigo, subgrupo_nombre
    grupo_seleccionado = grupo_combo.get()
    # Los subgrupos se actualizan con tanto el número como el nombre
    subgrupos = list(grupos[grupo_seleccionado])
    subgrupo_combo['values'] = subgrupos
    subgrupo_seleccionado = subgrupo_combo.get()  # Guardamos el subgrupo seleccionado

    # Verificamos si el subgrupo seleccionado es válido y lo separamos
    if subgrupo_seleccionado:
        subgrupo_codigo, subgrupo_nombre = separar_subgrupo(subgrupo_seleccionado)
        print(f"Subgrupo seleccionado: {subgrupo_codigo} - {subgrupo_nombre}")
        
        # Mostrar la separación en la interfaz gráfica
        label_separacion.config(text=f"CIU: {subgrupo_codigo} - Nombre actividad: {subgrupo_nombre}")

# Función para separar el código y el nombre del subgrupo
def separar_subgrupo(subgrupo):
    # Separar el código y el nombre del subgrupo
    if " - " in subgrupo:
        codigo, nombre = subgrupo.split(" - ", 1)  
        return codigo, nombre.strip()  
    else:
        return None, None

# Función para generar el documento Word con los datos del formulario
def generar_word():
    # Obtener la ruta del escritorio de forma dinámica
    ruta_escritorio = os.path.join(os.path.expanduser("~"), "Desktop")
    
    # Definir el archivo de plantilla según la opción seleccionada
    if combo_box.get() == "Aprobación exención contribución":
        archivo_plantilla = os.path.join(ruta_escritorio, "Plantilla_aprobación_exención.docx")
    else:  # Si es "Negación exención contribución"
        archivo_plantilla = os.path.join(ruta_escritorio, "Negación_Exención.docx")
    
    # Cargar el archivo de plantilla
    doc = Document(archivo_plantilla)
    
    # Verificar si el municipio está en la lista de municipios de Norte de Santander
    if municipio.get() in municipios_norte_santander:
        departamento = "Norte de Santander"
    else:
        departamento = "Cesar"
    
    # Reemplazar los marcadores con los datos del formulario en los párrafos
    for parrafo in doc.paragraphs:
        for run in parrafo.runs:
            # Reemplazo de 'CIU', 'ACTIVIDAD', 'DIRECCION', 'DEPARTAMENTO' con los valores correctos
            if '(CIU)' in run.text:
                run.text = run.text.replace('(CIU)', subgrupo_codigo)
            if '(ACTIVIDAD)' in run.text:
                run.text = run.text.replace('(ACTIVIDAD)', subgrupo_nombre)
            if '(DIRECCION)' in run.text:
                run.text = run.text.replace('(DIRECCION)', direccion.get())
            if '(DEPARTAMENTO)' in run.text:
                run.text = run.text.replace('(DEPARTAMENTO)', departamento)

            if 'NOMBRE_USUARIO' in run.text:
                run.text = run.text.replace('NOMBRE_USUARIO', nombre_usuario.get())
            if 'COD_CLIENTE' in run.text:
                run.text = run.text.replace('COD_CLIENTE', codigo_cliente.get())
            if 'CORREO' in run.text:
                run.text = run.text.replace('CORREO', correo.get())
            if 'CELULAR' in run.text:
                run.text = run.text.replace('CELULAR', telefono.get())
            if 'MUNICIPIO' in run.text:
                run.text = run.text.replace('MUNICIPIO', municipio.get())
            if 'RADICADO' in run.text:
                run.text = run.text.replace('RADICADO', radicado.get())
            if 'FECHA_RADICADO' in run.text:
                run.text = run.text.replace('FECHA_RADICADO', fecha_radicado.get())
            if 'EXPEDIENTE' in run.text:
                run.text = run.text.replace('EXPEDIENTE', expediente.get())
            if 'PROCESO' in run.text:
                run.text = run.text.replace('PROCESO', proceso.get())
            # Reemplazo para 'FECHA_PRUEBA'
            if 'FECHA_PRUEBA' in run.text:
                run.text = run.text.replace('FECHA_PRUEBA', fecha_radicado.get())
    
    # Reemplazar los marcadores dentro de las celdas de la tabla
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        # Reemplazo de los marcadores en las celdas
                        if '(CIU)' in run.text:
                            run.text = run.text.replace('(CIU)', subgrupo_codigo)
                        if '(ACTIVIDAD)' in run.text:
                            run.text = run.text.replace('(ACTIVIDAD)', subgrupo_nombre)
                        if '(DIRECCION)' in run.text:
                            run.text = run.text.replace('(DIRECCION)', direccion.get())
                        if '(DEPARTAMENTO)' in run.text:
                            run.text = run.text.replace('(DEPARTAMENTO)', departamento)

                        if 'NOMBRE_USUARIO' in run.text:
                            run.text = run.text.replace('NOMBRE_USUARIO', nombre_usuario.get())
                        if 'COD_CLIENTE' in run.text:
                            run.text = run.text.replace('COD_CLIENTE', codigo_cliente.get())
                        if 'CORREO' in run.text:
                            run.text = run.text.replace('CORREO', correo.get())
                        if 'CELULAR' in run.text:
                            run.text = run.text.replace('CELULAR', telefono.get())
                        if 'MUNICIPIO' in run.text:
                            run.text = run.text.replace('MUNICIPIO', municipio.get())
                        if 'RADICADO' in run.text:
                            run.text = run.text.replace('RADICADO', radicado.get())
                        if 'FECHA_RADICADO' in run.text:
                            run.text = run.text.replace('FECHA_RADICADO', fecha_radicado.get())
                        if 'EXPEDIENTE' in run.text:
                            run.text = run.text.replace('EXPEDIENTE', expediente.get())
                        if 'PROCESO' in run.text:
                            run.text = run.text.replace('PROCESO', proceso.get())
                        # Reemplazo para 'FECHA_PRUEBA'
                        if 'FECHA_PRUEBA' in run.text:
                            run.text = run.text.replace('FECHA_PRUEBA', fecha_radicado.get())

    # Guardar el documento con los cambios
    doc.save(os.path.join(ruta_escritorio, "Aprobación o Negación exención.docx"))

    print("Documento Word generado exitosamente.")

# Función para reiniciar el formulario cuando se cambie de opción
def reiniciar_formulario(event):
    mostrar_formulario()

# Función para resetear el formulario a su estado original
def resetear_formulario():
    combo_box.set("Seleccione una opción")  # Vuelve a poner la opción original
    # Limpiar el formulario y los desplegables
    for widget in frame_formulario.winfo_children():
        widget.destroy()
    for widget in frame_desplegables.winfo_children():
        widget.destroy()
    # Volver a crear el formulario con las opciones iniciales
    crear_campos_formulario()

# Menú desplegable de opciones
opciones = [
    "Aprobación exención contribución",
    "Negación exención contribución"
]

combo_box = ttk.Combobox(ventana, values=opciones, state="readonly")
combo_box.set("Seleccione una opción")
combo_box.pack(pady=10)

# Actualizar el formulario cada vez que se cambie la opción
combo_box.bind("<<ComboboxSelected>>", reiniciar_formulario)

# Botón para resetear el formulario
resetear_button = tk.Button(ventana, text="Resetear", command=resetear_formulario)
resetear_button.pack(pady=10)

# Botón para generar el Word
generar_button = tk.Button(ventana, text="Generar Word", command=generar_word)
generar_button.pack(pady=10)

# Etiqueta para mostrar la separación de código y nombre del subgrupo
label_separacion = tk.Label(ventana, text="Subgrupo Seleccionado:")
label_separacion.pack(pady=10)

# Botón para actualizar el subgrupo
actualizar_button = tk.Button(ventana, text="Actualizar", command=actualizar_subgrupos)
actualizar_button.pack(pady=10)

# Ejecutar la aplicación
ventana.mainloop()
