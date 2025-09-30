import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime
import csv
import os
from docx import Document

# =======================
# Variables "globales"
# =======================
nombre_usuario = ""
codigo_cliente = ""
direccion = ""
correo = ""
telefono = ""
municipio = ""
radicado = ""
fecha_radicado = ""
expediente = ""
proceso = ""

# Lista de municipios de Norte de Santander
municipios_norte_santander = [
    "Arboledas", "Cucutilla", "Gramalote", "Lourdes", "Salazar de Las Palmas",
    "Santiago", "Villa Caro", "Cúcuta", "El Zulia", "Los Patios",
    "Puerto Santander", "San Cayetano", "Villa del Rosario", "Bucarasica",
    "El Tarra", "Sardinata", "Tibú", "Ábrego", "Cáchira", "Convención",
    "El Carmen", "Hacarí", "La Esperanza", "La Playa de Belén", "Ocaña",
    "San Calixto", "Teorama", "Cácota", "Chitagá", "Mutiscua",
    "Pamplona", "Pamplonita", "Santo Domingo de Silos", "Bochalema",
    "Chinácota", "Durania", "Herrán", "Labateca", "Ragonvalia", "Toledo"
]

APP_TITLE = "Gestión de Trámites"
CSV_FILE = "tramites.csv"
CAMPOS = [
    "tipo_tramite",
    "nombre_usuario",
    "codigo_cliente",
    "direccion",
    "correo",
    "telefono",
    "municipio",
    "radicado",
    "fecha_radicado",
    "expediente",
    "proceso",
]

TIPOS_TRAMITE = [
    "Continuidad_provisional",       # Nombre exacto para el archivo
    "Reclamo_1_periodo",             # Nombre exacto para el archivo
    "Reclamo_mas_1_periodo",        # Nombre exacto para el archivo
]

def asegurar_csv():
    if not os.path.exists(CSV_FILE):
        with open(CSV_FILE, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=CAMPOS)
            writer.writeheader()

def obtener_ruta_escritorio():
    if os.name == 'nt':  # Windows
        # Ruta del escritorio en Windows
        return os.path.join(os.environ['USERPROFILE'], "Desktop")
    else:  # Linux/macOS
        # Ruta del escritorio en Linux/macOS
        return os.path.expanduser("~/Escritorio")
    
def buscar_documento(tipo_tramite):
    escritorio = obtener_ruta_escritorio()  # Llamamos a la función que determina la ruta del escritorio
    
    print(f"Buscando en el escritorio: {escritorio}")  # Verificación de la ruta

    if not os.path.exists(escritorio):
        print("No se encontró la ruta del escritorio")
        return None
    
    if tipo_tramite == "Continuidad_provisional":
        documento = os.path.join(escritorio, "Continuidad_provisional.docx")
    elif tipo_tramite == "Reclamo_1_periodo":
        documento = os.path.join(escritorio, "Reclamo_1_periodo.docx")
    elif tipo_tramite == "Reclamo_mas_1_periodo":
        documento = os.path.join(escritorio, "Reclamo_mas_1_periodo.docx")
    else:
        return None
    
    # Verificamos si el archivo existe y si lo encontramos
    if os.path.exists(documento):
        print(f"Documento encontrado: {documento}")
        return documento
    else:
        print(f"No se encontró el archivo: {documento}")
        return None

def reemplazar_marcadores(doc, datos):
    # Reemplazar los marcadores con los datos del formulario en los párrafos
    for parrafo in doc.paragraphs:
        for run in parrafo.runs:
            for key, value in datos.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)
    
    # Reemplazar los marcadores dentro de las celdas de la tabla
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        for key, value in datos.items():
                            if key in run.text:
                                run.text = run.text.replace(key, value)
    return doc

class FormController:
    def __init__(self, parent):
        self.parent = parent
        self.frame = ttk.Frame(parent)
        self.frame.columnconfigure(0, weight=1)
        self.frame.columnconfigure(1, weight=2)

        # Estado del formulario
        self.tipo_tramite = tk.StringVar(value="")
        self.vars = {
            "nombre_usuario": tk.StringVar(),
            "codigo_cliente": tk.StringVar(),
            "direccion": tk.StringVar(),
            "correo": tk.StringVar(),
            "telefono": tk.StringVar(),
            "municipio": tk.StringVar(),
            "radicado": tk.StringVar(value=""),        # VACÍO
            "fecha_radicado": tk.StringVar(value=""),  # VACÍO
            "expediente": tk.StringVar(),
            "proceso": tk.StringVar(),
        }

        # Encabezado
        self.lbl_title = ttk.Label(
            self.frame, text="Complete el formulario", style="Title.TLabel"
        )
        self.lbl_title.grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 8))

        # Campos
        self._row = 1
        self.build_entry("Nombre del usuario", "nombre_usuario", required=True)
        self.build_entry("Código de cliente", "codigo_cliente", required=True)
        self.build_entry("Dirección", "direccion")
        self.build_entry("Correo", "correo")
        self.build_entry("Teléfono", "telefono", required=True)
        self.build_entry("Municipio", "municipio", required=True)
        self.build_entry("Radicado", "radicado")
        self.build_entry("Fecha de radicado (DD/MM/YYYY)", "fecha_radicado")
        self.build_entry("Expediente", "expediente")
        self.build_entry("Proceso", "proceso")

        # Botones
        btns = ttk.Frame(self.frame)
        btns.grid(row=self._row, column=0, columnspan=2, pady=(10, 0), sticky="e")
        self._row += 1

        self.btn_limpiar = ttk.Button(btns, text="Limpiar", command=self.on_limpiar)
        self.btn_limpiar.grid(row=0, column=0, padx=6)

        self.btn_guardar = ttk.Button(btns, text="Generar Documento", command=self.on_guardar)
        self.btn_guardar.grid(row=0, column=1, padx=6)

        # Status
        self.status = ttk.Label(self.frame, text="", style="Muted.TLabel")
        self.status.grid(row=self._row, column=0, columnspan=2, sticky="w", pady=(8,0))

    def grid(self, **kwargs):
        self.frame.grid(**kwargs)

    def set_tipo_tramite(self, tipo):
        self.tipo_tramite.set(tipo)
        self.lbl_title.configure(text=f"Formulario — {tipo}" if tipo else "Complete el formulario")
        self.status.configure(text=f"Trámite seleccionado: {tipo}" if tipo else "")

    def build_entry(self, label, key, required=False):
        lbl = ttk.Label(self.frame, text=f"{label}{' *' if required else ''}")
        lbl.grid(row=self._row, column=0, sticky="w", padx=(0, 12), pady=3)

        ent = ttk.Entry(self.frame, textvariable=self.vars[key])
        ent.grid(row=self._row, column=1, sticky="ew", pady=3)

        self._row += 1

    def on_limpiar(self):
        for v in self.vars.values():
            v.set("")
        self.status.configure(text="Formulario limpiado.")

    def validar(self):
        requeridos = ["nombre_usuario", "codigo_cliente", "telefono", "municipio"]
        vacios = [k for k in requeridos if not self.vars[k].get().strip()]
        if vacios:
            campos = ", ".join(vacios)
            messagebox.showwarning(
                "Campos requeridos",
                f"Por favor complete los campos obligatorios: {campos}",
            )
            return False

        # Validación de fecha SOLO si hay algo escrito
        fecha = self.vars["fecha_radicado"].get().strip()
        if fecha:
            try:
                # Cambiamos el formato de la fecha a DD/MM/YYYY
                datetime.strptime(fecha, "%d/%m/%Y")
            except ValueError:
                messagebox.showwarning(
                    "Fecha inválida",
                    "Ingrese la fecha de radicado con formato DD/MM/YYYY.",
                )
                return False

        return True

    def on_guardar(self):
        if not self.tipo_tramite.get():
            messagebox.showinfo("Seleccione trámite", "Seleccione un tipo de trámite.")
            return
        if not self.validar():
            return

        # Verificar si el municipio está en la lista de municipios de Norte de Santander
        if self.vars["municipio"].get() in municipios_norte_santander:
            departamento = "Norte de Santander"
        else:
            departamento = "Cesar"

        # Obtener el documento correspondiente según el tipo de trámite
        documento = buscar_documento(self.tipo_tramite.get())
        
        if not documento:
            messagebox.showerror("Error", "No se encontró el documento para el trámite seleccionado.")
            return

        # Reemplazar los marcadores en el documento
        doc = Document(documento)

        # Asegurarnos de convertir la fecha al formato DD/MM/YYYY
        fecha_radicado_formateada = self.vars['fecha_radicado'].get()

        datos = {
            '(DIRECCION)': self.vars['direccion'].get(),
            '(DEPARTAMENTO)': departamento,  # Usamos la variable departamento
            'NOMBRE_USUARIO': self.vars['nombre_usuario'].get(),
            'COD_CLIENTE': self.vars['codigo_cliente'].get(),
            'CORREO': self.vars['correo'].get(),
            'CELULAR': self.vars['telefono'].get(),
            'MUNICIPIO': self.vars['municipio'].get(),
            'RADICADO': self.vars['radicado'].get(),
            'FECHA': fecha_radicado_formateada,  # Usamos la fecha en el nuevo formato
            'EXPEDIENTE': self.vars['expediente'].get(),
            'PROCESO': self.vars['proceso'].get(),
            'FECHA_PRUEBA': fecha_radicado_formateada
        }

        doc = reemplazar_marcadores(doc, datos)

        # Guardar el nuevo documento en el escritorio
        ruta_guardado = os.path.join(obtener_ruta_escritorio(), f"Tramite_{self.tipo_tramite.get()}.docx")
        doc.save(ruta_guardado)

        messagebox.showinfo("Documento Generado", f"Documento generado con éxito en: {ruta_guardado}")

class App:
    def __init__(self, root):
        self.root = root
        self.root.title(APP_TITLE)
        self.root.geometry("900x520")
        self.root.minsize(820, 480)

        style = ttk.Style()
        try:
            style.theme_use("clam")
        except:
            pass
        style.configure("Title.TLabel", font=("Segoe UI", 14, "bold"))
        style.configure("Muted.TLabel", foreground="#666666")
        style.configure("Card.TFrame", relief="groove", borderwidth=1, padding=12)

        self.root.columnconfigure(0, weight=1, minsize=280)
        self.root.columnconfigure(1, weight=2)
        self.root.rowconfigure(0, weight=1)

        self.panel_opciones = ttk.Frame(self.root, padding=12)
        self.panel_opciones.grid(row=0, column=0, sticky="nsew")
        self.panel_opciones.columnconfigure(0, weight=1)

        self.form_container = ttk.Frame(self.root, padding=12)
        self.form_container.grid(row=0, column=1, sticky="nsew")
        self.form_container.columnconfigure(0, weight=1)

        titulo = ttk.Label(
            self.panel_opciones,
            text="Seleccione el tipo de trámite",
            style="Title.TLabel",
        )
        titulo.grid(row=0, column=0, sticky="w", pady=(0, 8))

        for i, texto in enumerate(TIPOS_TRAMITE, start=1):
            card = ttk.Frame(self.panel_opciones, style="Card.TFrame")
            card.grid(row=i, column=0, sticky="ew", pady=6)
            card.columnconfigure(0, weight=1)

            lbl = ttk.Label(card, text=texto, font=("Segoe UI", 11, "bold"))
            lbl.grid(row=0, column=0, sticky="w")
            desc = ttk.Label(
                card,
                text="Haga clic para diligenciar el formulario.",
                style="Muted.TLabel",
            )
            desc.grid(row=1, column=0, sticky="w", pady=(2, 0))

            btn = ttk.Button(card, text="Seleccionar", command=lambda t=texto: self.on_select(t))
            btn.grid(row=0, column=1, rowspan=2, padx=(12, 0))

        self.form = FormController(self.form_container)
        self.form.grid(row=0, column=0, sticky="nsew")

        self.form.set_tipo_tramite("")

    def on_select(self, tipo):
        self.form.set_tipo_tramite(tipo)

def main():
    asegurar_csv()
    root = tk.Tk()
    App(root)
    root.mainloop()

if __name__ == "__main__":
    main()
